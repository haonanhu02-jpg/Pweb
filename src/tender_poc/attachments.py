from __future__ import annotations

import re
import shutil
import subprocess
import zipfile
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import unquote, urlparse

import requests
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
import xlrd

from tender_poc.models import AttachmentDocumentV1, TenderNoticeV1, make_content_hash, normalize_text, now_utc
from tender_poc.paths import ATTACHMENTS_DIR, ensure_data_dirs
from tender_poc.storage import TenderStore


SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".zip"}
UNSUPPORTED_EXTENSIONS = {".rar"}
MAX_DOWNLOAD_BYTES = 25 * 1024 * 1024
MAX_TEXT_CHARS = 200_000


@dataclass
class AttachmentParseSummary:
    parsed_count: int = 0
    skipped_count: int = 0
    empty_count: int = 0
    unsupported_count: int = 0
    missing_tool_count: int = 0
    failed_count: int = 0

    def as_dict(self) -> dict[str, int]:
        return {
            "parsed_count": self.parsed_count,
            "skipped_count": self.skipped_count,
            "empty_count": self.empty_count,
            "unsupported_count": self.unsupported_count,
            "missing_tool_count": self.missing_tool_count,
            "failed_count": self.failed_count,
        }


def parse_notice_attachments(
    store: TenderStore,
    notices: list[TenderNoticeV1],
    *,
    limit: int | None = None,
    timeout_seconds: int = 30,
) -> AttachmentParseSummary:
    ensure_data_dirs()
    parser = AttachmentParser(store=store, timeout_seconds=timeout_seconds)
    summary = AttachmentParseSummary()
    processed = 0
    for notice in notices:
        for attachment in notice.attachments:
            if limit is not None and processed >= limit:
                return summary
            processed += 1
            status = parser.parse_one(
                notice_id=notice.id,
                attachment_name=attachment.name,
                attachment_url=attachment.url,
            )
            _increment_summary(summary, status)
    return summary


class AttachmentParser:
    def __init__(self, store: TenderStore, timeout_seconds: int = 30) -> None:
        self.store = store
        self.timeout_seconds = timeout_seconds
        self.session = requests.Session()
        self.session.headers.update(
            {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/126.0 Safari/537.36 tender-poc/0.1"
                )
            }
        )

    def parse_one(self, *, notice_id: str, attachment_name: str, attachment_url: str) -> str:
        existing = self.store.get_attachment_document(notice_id, attachment_url)
        if existing and existing["status"] in {"parsed", "empty"}:
            return "skipped"

        ext = _detect_extension(attachment_name, attachment_url)
        if ext in UNSUPPORTED_EXTENSIONS:
            self._save_document(
                notice_id=notice_id,
                attachment_url=attachment_url,
                attachment_name=attachment_name,
                file_ext=ext,
                status="unsupported",
                error=f"Unsupported attachment type: {ext}",
            )
            return "unsupported"
        if ext and ext not in SUPPORTED_EXTENSIONS:
            self._save_document(
                notice_id=notice_id,
                attachment_url=attachment_url,
                attachment_name=attachment_name,
                file_ext=ext,
                status="unsupported",
                error=f"Unsupported attachment type: {ext}",
            )
            return "unsupported"

        try:
            file_path = self._download(notice_id, attachment_name, attachment_url, ext)
            ext = file_path.suffix.lower()
            text = _extract_text(file_path)
            text = normalize_text(text)[:MAX_TEXT_CHARS]
            status = "parsed" if text else "empty"
            self._save_document(
                notice_id=notice_id,
                attachment_url=attachment_url,
                attachment_name=attachment_name,
                file_path=file_path,
                file_ext=ext,
                status=status,
                content_text=text,
                content_hash=make_content_hash(text) if text else None,
                fetched_at=now_utc(),
            )
            return status
        except MissingToolError as exc:
            self._save_document(
                notice_id=notice_id,
                attachment_url=attachment_url,
                attachment_name=attachment_name,
                file_ext=ext,
                status="missing_tool",
                error=str(exc),
            )
            return "missing_tool"
        except UnsupportedAttachmentError as exc:
            self._save_document(
                notice_id=notice_id,
                attachment_url=attachment_url,
                attachment_name=attachment_name,
                file_ext=ext,
                status="unsupported",
                error=str(exc),
            )
            return "unsupported"
        except Exception as exc:  # noqa: BLE001 - attachment failures must not abort collection.
            self._save_document(
                notice_id=notice_id,
                attachment_url=attachment_url,
                attachment_name=attachment_name,
                file_ext=ext,
                status="failed",
                error=repr(exc),
            )
            return "failed"

    def _download(self, notice_id: str, attachment_name: str, attachment_url: str, ext: str | None) -> Path:
        response = self.session.get(attachment_url, timeout=self.timeout_seconds)
        response.raise_for_status()
        if len(response.content) > MAX_DOWNLOAD_BYTES:
            raise ValueError(f"Attachment exceeds {MAX_DOWNLOAD_BYTES} bytes")

        ext = ext or _extension_from_content_type(response.headers.get("content-type")) or ".bin"
        directory = ATTACHMENTS_DIR / notice_id
        directory.mkdir(parents=True, exist_ok=True)
        path = directory / _safe_filename(attachment_name, attachment_url, ext)
        path.write_bytes(response.content)
        return path

    def _save_document(
        self,
        *,
        notice_id: str,
        attachment_url: str,
        attachment_name: str,
        file_ext: str | None,
        status: str,
        file_path: Path | None = None,
        content_text: str = "",
        content_hash: str | None = None,
        error: str | None = None,
        fetched_at=None,
    ) -> None:
        self.store.upsert_attachment_document(
            AttachmentDocumentV1(
                notice_id=notice_id,
                attachment_url=attachment_url,
                attachment_name=attachment_name,
                file_path=str(file_path) if file_path else None,
                file_ext=file_ext,
                status=status,
                content_text=content_text,
                content_hash=content_hash,
                error=error,
                fetched_at=fetched_at,
            )
        )


class MissingToolError(RuntimeError):
    pass


class UnsupportedAttachmentError(RuntimeError):
    pass


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == ".doc":
        return _extract_doc(path)
    if ext == ".xlsx":
        return _extract_xlsx(path)
    if ext == ".xls":
        return _extract_xls(path)
    if ext == ".zip":
        return _extract_zip(path)
    raise UnsupportedAttachmentError(f"Unsupported attachment type: {ext or 'unknown'}")


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_doc(path: Path) -> str:
    antiword = shutil.which("antiword")
    if not antiword:
        raise MissingToolError("antiword is required to parse .doc files")
    completed = subprocess.run(
        [antiword, str(path)],
        check=True,
        capture_output=True,
        text=True,
        timeout=30,
    )
    return completed.stdout


def _extract_xlsx(path: Path) -> str:
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        return _workbook_cells_to_text(workbook.worksheets)
    finally:
        workbook.close()


def _extract_xls(path: Path) -> str:
    workbook = xlrd.open_workbook(str(path))
    parts = []
    for sheet in workbook.sheets():
        parts.append(f"Sheet: {sheet.name}")
        for row_index in range(sheet.nrows):
            values = [str(value) for value in sheet.row_values(row_index) if value not in ("", None)]
            if values:
                parts.append(" ".join(values))
    return "\n".join(parts)


def _extract_zip(path: Path) -> str:
    target_dir = path.with_suffix("")
    target_dir.mkdir(parents=True, exist_ok=True)
    parts = []
    with zipfile.ZipFile(path) as archive:
        for member in archive.infolist():
            if member.is_dir():
                continue
            inner_ext = Path(member.filename).suffix.lower()
            if inner_ext not in SUPPORTED_EXTENSIONS - {".zip"}:
                continue
            safe_name = _safe_inner_filename(member.filename)
            inner_path = target_dir / safe_name
            inner_path.write_bytes(archive.read(member))
            try:
                text = normalize_text(_extract_text(inner_path))
            except MissingToolError:
                raise
            except Exception:
                continue
            if text:
                parts.append(f"{safe_name}\n{text}")
    return "\n\n".join(parts)


def _workbook_cells_to_text(worksheets) -> str:
    parts = []
    for worksheet in worksheets:
        parts.append(f"Sheet: {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            values = [str(value) for value in row if value not in ("", None)]
            if values:
                parts.append(" ".join(values))
    return "\n".join(parts)


def _detect_extension(name: str, url: str) -> str | None:
    for value in (urlparse(url).path, unquote(urlparse(url).path), name):
        ext = Path(value).suffix.lower()
        if ext:
            return ext
    return None


def _extension_from_content_type(content_type: str | None) -> str | None:
    if not content_type:
        return None
    content_type = content_type.lower()
    if "pdf" in content_type:
        return ".pdf"
    if "spreadsheet" in content_type or "excel" in content_type:
        return ".xlsx"
    if "word" in content_type:
        return ".docx"
    if "zip" in content_type:
        return ".zip"
    return None


def _safe_filename(name: str, url: str, ext: str) -> str:
    raw_stem = Path(unquote(urlparse(url).path)).stem or Path(name).stem or "attachment"
    stem = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "_", raw_stem).strip("._")
    if not stem:
        stem = "attachment"
    suffix = ext if ext.startswith(".") else f".{ext}"
    return f"{stem[:90]}{suffix}"


def _safe_inner_filename(name: str) -> str:
    raw_name = Path(name).name or "inner_attachment"
    return re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "_", raw_name)


def _increment_summary(summary: AttachmentParseSummary, status: str) -> None:
    if status == "parsed":
        summary.parsed_count += 1
    elif status == "skipped":
        summary.skipped_count += 1
    elif status == "empty":
        summary.empty_count += 1
    elif status == "unsupported":
        summary.unsupported_count += 1
    elif status == "missing_tool":
        summary.missing_tool_count += 1
    else:
        summary.failed_count += 1
